#!/usr/bin/env python3
"""
Extraction functions for each file type to be used with Document class

Notes on getting text:
* assume docs are max 100 pages - only get N bytes of string to conserve memory
* `text_lst` itesms should be in batches, such as doc pages or paragraphs (don't mess up sentenceizer)
* `docs = spacy.pipe(text_lst, n_process=-1)`   #as many processes as CPUs
* resulting docs is a generator, but return as list `list(docs)`, currently
* then Document performs text processing

TODO:
* convert all file types to pdf, then extract using pdf.miner
* limit retrieval to N excerpts
* use optimal storage: `from collections import namedtuple`   #ref: https://stackoverflow.com/questions/1336791/dictionary-vs-object-which-is-more-efficient-and-why
"""

__author__ = "Jason Beach"
__version__ = "0.1.0"
__license__ = "MIT"



import pandas as pd

#docx
import docx

#html
import bs4
from bs4.element import Comment
from xhtml2pdf import pisa 

#pdf
import pypdf
import pdftitle                                                                 #uses pdfminer
from pdfminer.high_level import extract_text as pdf_extract_text                #uses pdfminer.six, problem TODO???
from pdfminer.high_level import extract_pages as pdf_extract_pages
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
import fitz


from utils import record




def get_title(txt):
    pass

def clean_text(txt):
    if type(txt) is list:
        combined_txt = ('.').join(txt)
        return combined_txt.replace('.','.  ').replace('\n',' ')
    elif type(txt) is str:
        txts = txt.split('.\n')
        txts = [txt.replace('-\n','').replace('\n',' ') for txt in txts]
        return txts
    else:
        return txt

def html_string_to_pdf(content, output):
    """
    Generate a pdf using a string content

    content : str - content to write in the pdf file
    output  : str - name of the file to create
    """
    # Open file to write
    result_file = open(output, "w+b") # w+b to write in binary mode.
    # convert HTML to PDF
    pisa_status = pisa.CreatePDF(
            content,                   # the HTML to convert
            dest=result_file           # file handle to recieve result
    )           
    # close output file
    result_file.close()
    result = pisa_status.err
    if not result:
        print("Successfully created PDF")
    else:
        print("Error: unable to create the PDF")    
    # return False on success and True on errors
    return result







def extract_docx_original(self, logger):
    """Extract from docx filetype.
    
    'text': excerpts.paragraphs
    """
    try:
        excerpts = docx.Document(self.filepath)
    except:
        logger.info("TypeError: document not of type `.docx`")
        return None
    txt = [par.text for par in excerpts.paragraphs]

    record['title'] = excerpts.paragraphs[0].text     #<<< get_title(txt)
    record['length_lines'] = len( clean_text(txt).split('.') )
    record['text'] = txt
    return record


def extract_docx(self, logger):
    """Extract from docx using `extract_pdf`
    
    'text': excerpts.paragraphs
    """
    try:
        excerpts = docx.Document(self.filepath)
    except:
        logger.info("TypeError: document not of type `.docx`")
        return None
    txt_lst = [par.text for par in excerpts.paragraphs]
    txt = ''.join(txt_lst)

    tmp_file = self.filepath.parent / (self.filepath.stem + '.pdf')
    html_string_to_pdf(content = txt, 
                       output = tmp_file
                      )
    self.filepath = tmp_file
    record = extract_pdf(self, logger)
    return record


def extract_html_original(self, logger):
    """Extract for html filetype.
    
    'text': visible text
    """
    #ref: https://stackoverflow.com/questions/1936466/how-to-scrape-only-visible-webpage-text-with-beautifulsoup
    def tag_visible(element):
        if element.parent.name in ['style', 'script', 'head', 'title', 'meta', '[document]']:
            return False
        if isinstance(element, Comment):
            return False
        return True

    def text_from_html(soup):
        texts = soup.findAll(text=True)
        visible_texts = filter(tag_visible, texts)  
        return u" ".join(t.strip() for t in visible_texts)

    with open(self.filepath.__str__(), 'r') as f:
        html = f.readlines()
    clean_html = ('').join(html).replace('\n','')
    try:
        excerpts = bs4.BeautifulSoup(clean_html, 'html.parser')
    except:
        logger.info("TypeError: document not of type `.html`")
        return None
    txt = text_from_html(excerpts)
    record['title'] = excerpts.find('title').text
    record['text'] = txt
    return record
    

def extract_html(self, logger):
    """Extract from html using `extract_pdf`
    TODO: fix so that .pdf is removed when finished
    """
    with open(self.filepath.__str__(), 'r') as f:
        html = f.read()
    tmp_file = self.filepath.parent / (self.filepath.stem + '.pdf')
    html_string_to_pdf(content = html, 
                       output = tmp_file
                      )
    self.filepath = tmp_file
    record = extract_pdf(self, logger)
    return record


def extract_pdf(self, logger):
    """Extract from pdf filetype.
    
    'text': pages
    """
    def get_title(title):
        if not title:
            with open(self.filepath.__str__(), 'rb') as f:
                pdfReader = pypdf.PdfReader(f)
                first_str = len(excerpts[0])
                if pdfReader.metadata['/Title']:
                    title = pdfReader.metadata['/Title']
                elif first_str>0 and first_str<100:
                    title = first_str
                else:
                    pass
        return title

    def get_toc():
        outlines = None
        doc = fitz.open(self.filepath.__str__())
        tmp = doc.get_toc()
        outlines = tmp if tmp != [] else None
        if not outlines:
            with open(self.filepath.__str__(), 'rb') as fp:
                parser = PDFParser(fp)
                document = PDFDocument(parser)
                try:
                    outlines = list(document.get_outlines())
                except:
                    pass
        return outlines


    title = None
    try:
        title = pdftitle.get_title_from_file(self.filepath.__str__())
    except Exception:
        logger.info("`pdftitle` module threw error")
        pass
    raw_text = pdf_extract_text(self.filepath.__str__())
    excerpts = clean_text(raw_text)
    pages_gen = pdf_extract_pages(self.filepath.__str__())
    page_nos = sum(1 for x in pages_gen)

    record['title'] = get_title(title)
    record['page_nos'] = page_nos
    record['toc'] = get_toc()
    record['text'] = excerpts

    return record


def extract_csv(self, logger):
    df = pd.read_csv(self.filepath.__str__(), 
                    nrows=1
                    )
    record['title'] = df.columns[0].replace('\n','')
    return record


def extract_xlsx(self, logger):
    sheets = pd.ExcelFile(self.filepath.__str__()).sheet_names
    df = pd.read_excel(self.filepath.__str__(), 
                        sheet_name=sheets[0], 
                        nrows=1
                        )

    record['title'] = df.columns[0].replace('\n','')
    return record